import streamlit as st
from pathlib import Path
import numpy as np
from PIL import Image, ImageDraw, ImageFont
import io
import base64
from gtts import gTTS
import tempfile
import regex as re
import chardet
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import groq

class BionicTextConverter:
    def __init__(self):
        self.default_settings = {
            'emphasis_level': 0.5,
            'emphasis_color': '#FF0000',
            'background_color': '#FFFFFF',
            'font_size': 24,
            'language': 'Hindi'
        }
        
        self.supported_languages = {
            'Hindi': {'name': 'Hindi', 'code': 'hi', 'placeholder': 'यहाँ हिंदी टेक्स्ट दर्ज करें...', 'help': 'Press Ctrl/Cmd + Enter to convert text'},
            'Marathi': {'name': 'Marathi', 'code': 'mr', 'placeholder': 'येथे मराठी मजकूर टाइप करा...', 'help': 'Press Ctrl/Cmd + Enter to convert text'},
            'English': {'name': 'English', 'code': 'en', 'placeholder': 'Enter text here...', 'help': 'Press Ctrl/Cmd + Enter to convert text'},
            'German': {'name': 'Deutsch', 'code': 'de', 'placeholder': 'Text hier eingeben...', 'help': 'Drücken Sie Strg/Cmd + Enter zum Konvertieren'},
            'French': {'name': 'Français', 'code': 'fr', 'placeholder': 'Entrez le texte ici...', 'help': 'Appuyez sur Ctrl/Cmd + Entrée pour convertir'},
            'Spanish': {'name': 'Español', 'code': 'es', 'placeholder': 'Ingrese el texto aquí...', 'help': 'Presione Ctrl/Cmd + Enter para convertir'},
            'Italian': {'name': 'Italiano', 'code': 'it', 'placeholder': 'Inserisci il testo qui...', 'help': 'Premi Ctrl/Cmd + Invio per convertire'}
        }
    
    def process_hindi_text(self, text: str, settings: dict) -> str:
        """Process Devanagari text (Hindi/Marathi) using grapheme clusters"""
        words = text.split()
        processed_words = []
        
        for word in words:
            graphemes = re.findall(r'\X', word)
            chars_to_emphasize = max(1, int(len(graphemes) * settings['emphasis_level']))
            emphasized = ''.join(graphemes[:chars_to_emphasize])
            normal = ''.join(graphemes[chars_to_emphasize:])
            processed_words.append(
                f"<span style='color: {settings['emphasis_color']}; font-weight: 1000'>{emphasized}</span>{normal}"
            )
        
        return " ".join(processed_words)
    
    def process_latin_text(self, text: str, settings: dict) -> str:
        """Process text in Latin-based scripts"""
        words = text.split()
        processed_words = []
        
        for word in words:
            if len(word) <= 3:
                chars_to_emphasize = 1
            else:
                chars_to_emphasize = max(2, int(len(word) * settings['emphasis_level']))
            
            emphasized = word[:chars_to_emphasize]
            normal = word[chars_to_emphasize:]
            
            processed_words.append(
                f"<span style='color: {settings['emphasis_color']}; font-weight: 1000'>{emphasized}</span>{normal}"
            )
        
        return " ".join(processed_words)
    
    def process_text(self, text: str, settings: dict) -> str:
        """Process text based on selected language"""
        if settings.get('language') in ['Hindi', 'Marathi']:  # Updated to handle both Devanagari scripts
            return self.process_hindi_text(text, settings)
        else:
            return self.process_latin_text(text, settings)
    
    def text_to_speech(self, text: str, language_code: str) -> tuple[bytes, str]:
        """Convert text to speech using gTTS"""
        try:
            # Generate a unique temporary file name
            temp_path = tempfile.mktemp(suffix='.mp3')
            
            # Create gTTS object and save directly to the temporary file
            tts = gTTS(text=text, lang=language_code, slow=False)
            tts.save(temp_path)
            
            # Read the audio file
            with open(temp_path, 'rb') as audio_file:
                audio_bytes = audio_file.read()
            
            return audio_bytes, "audio/mp3"
            
        except Exception as e:
            st.error(f"Error generating speech: {str(e)}")
            return None, None
            
        finally:
            # Clean up the temporary file
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass

    def export_to_word(self, text: str, settings: dict) -> bytes:
        """Export bionic text to Word document maintaining the same styling"""
        doc = Document()
        
        # Set font size for the document
        style = doc.styles['Normal']
        style.font.size = Pt(settings['font_size'])
        
        # Convert color from hex to RGB
        color_hex = settings['emphasis_color'].lstrip('#')
        rgb_emphasis = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        
        # Process text based on language
        if settings.get('language') in ['Hindi', 'Marathi']:
            words = text.split()
            para = doc.add_paragraph()
            for word in words:
                graphemes = re.findall(r'\X', word)
                chars_to_emphasize = max(1, int(len(graphemes) * settings['emphasis_level']))
                
                # Add emphasized part
                run = para.add_run(''.join(graphemes[:chars_to_emphasize]))
                run.font.bold = True
                run.font.color.rgb = RGBColor(*rgb_emphasis)
                
                # Add normal part
                run = para.add_run(''.join(graphemes[chars_to_emphasize:]))
                
                # Add space between words
                para.add_run(' ')
        else:
            words = text.split()
            para = doc.add_paragraph()
            for word in words:
                if len(word) <= 3:
                    chars_to_emphasize = 1
                else:
                    chars_to_emphasize = max(2, int(len(word) * settings['emphasis_level']))
                
                # Add emphasized part
                run = para.add_run(word[:chars_to_emphasize])
                run.font.bold = True
                run.font.color.rgb = RGBColor(*rgb_emphasis)
                
                # Add normal part
                run = para.add_run(word[chars_to_emphasize:])
                
                # Add space between words
                para.add_run(' ')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

    def summarize_text(self, text: str, language: str) -> str:
        """Summarize text using Groq API"""
        try:
            # Get API key from Streamlit secrets or use environment variable as fallback
            api_key = st.secrets.groq.api_key if hasattr(st.secrets, "groq") else "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"
            
            # Initialize Groq client
            groq_client = groq.Groq(api_key=api_key)

            # Prepare the prompt based on language
            language_name = self.supported_languages[language]['name']
            prompt = f"""Please provide a concise summary of the following {language_name} text. 
            Keep the summary in the same language as the input text.
            Focus on the main points and key information:

            {text}"""

            # Get completion from Groq
            completion = groq_client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional {language_name} language expert and summarizer."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                model="qwen-qwq-32b",
                temperature=0.6,
                max_tokens=4096
            )

            # Return the summary
            return completion.choices[0].message.content.strip()

        except Exception as e:
            st.error(f"Error generating summary: {str(e)}")
            return None

def main():
    st.title("🔍 Multi-Language Bionic Text Converter", anchor=False)
    st.caption("Convert text into a more readable format with visual emphasis and audio support")
    
    # Initialize converter
    converter = BionicTextConverter()
    
    # Sidebar for settings
    with st.sidebar:
        st.title("⚙️ Settings", anchor=False)
        st.markdown("---")
        
        # Language selection with info
        st.subheader("Language Settings", anchor=False)
        selected_language = st.selectbox(
            "Choose Language",
            list(converter.supported_languages.keys()),
            index=0,
            help="Select the language of your input text"
        )
        
        st.markdown("---")
        st.subheader("Visual Settings", anchor=False)
        
        # Theme and color settings
        emphasis_color = st.color_picker(
            "Emphasis Color",
            converter.default_settings['emphasis_color'],
            help="Choose the color for emphasized text parts"
        )
        
        theme = st.selectbox(
            "Theme",
            ["Light", "Dark"],
            help="Choose between light and dark theme"
        )
        
        # Text settings
        emphasis_level = st.slider(
            "Emphasis Level",
            0.0, 1.0,
            converter.default_settings['emphasis_level'],
            help="Control how much of each word is emphasized"
        )
        
        font_size = st.slider(
            "Font Size",
            12, 48,
            converter.default_settings['font_size'],
            help="Adjust the text size for better readability"
        )
    
    # Input methods tabs
    text_tab, file_tab = st.tabs(["✏️ Text Input", "📄 File Upload"])
    
    # Common settings dictionary
    settings = {
        'emphasis_level': emphasis_level,
        'emphasis_color': emphasis_color,
        'background_color': '#1E1E1E' if theme == "Dark" else '#FFFFFF',
        'font_size': font_size,
        'language': selected_language
    }
    
    with text_tab:
        # Text input area with keyboard shortcut hint
        input_text = st.text_area(
            f"Enter {selected_language} Text",
            height=200,
            placeholder=converter.supported_languages[selected_language]['placeholder'],
            help=converter.supported_languages[selected_language]['help'],
            key="text_input"
        )
        
        if input_text:
            # Add Summary section before Bionic Text
            st.markdown("---")
            st.markdown("### 📋 Summary")
            summary_col1, summary_col2 = st.columns([1, 3])
            
            with summary_col1:
                if st.button("📝 Generate Summary", use_container_width=True, help="Generate a summary of the text"):
                    with st.spinner("Generating summary..."):
                        summary = converter.summarize_text(input_text, selected_language)
                        if summary:
                            st.session_state.text_summary = summary
                            st.success("✨ Summary generated!")
            
            if 'text_summary' in st.session_state:
                with summary_col2:
                    st.markdown(
                        f"""
                        <div style='background-color: {settings["background_color"]}; 
                                  padding: 15px; 
                                  border-radius: 10px;
                                  font-size: {max(12, settings["font_size"]-4)}px;
                                  color: {"white" if theme == "Dark" else "black"};
                                  box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                                  margin: 5px 0;'>
                            {st.session_state.text_summary}
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

            st.markdown("---")
            st.markdown("### 📝 Bionic Text")
            # Process text
            processed_text = converter.process_text(input_text, settings)
            
            # Display processed text
            st.markdown(
                f"""
                <div style='background-color: {settings["background_color"]}; 
                          padding: 20px; 
                          border-radius: 10px;
                          font-size: {settings["font_size"]}px;
                          color: {"white" if theme == "Dark" else "black"};
                          box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                          margin: 10px 0;'>
                    {processed_text}
                </div>
                """,
                unsafe_allow_html=True
            )
            
            # Export to Word option
            st.markdown("---")
            st.markdown("### 📄 Export Options")
            export_col1, export_col2, export_col3 = st.columns([2, 4, 2])
            with export_col1:
                if st.button("📎 Export to Word", use_container_width=True, help="Download as Word document (.docx)"):
                    try:
                        doc_bytes = converter.export_to_word(input_text, settings)
                        with export_col2:
                            st.download_button(
                                label="⬇️ Download Word Document",
                                data=doc_bytes,
                                file_name=f"bionic_text_{settings['language'].lower()}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"Error creating Word document: {str(e)}")
            
            # Text to Speech section with improved layout
            st.markdown("---")
            st.markdown("### 🔊 Text to Speech")
            st.caption("Listen to the original text in its native language")
            
            cols = st.columns([2, 4, 2])
            with cols[0]:
                tts_button = st.button(
                    "🎙️ Generate Audio",
                    key="tts_text",
                    use_container_width=True,
                    help="Press Space or Enter when focused to generate audio"
                )
            
            with cols[1]:
                if tts_button:
                    with st.spinner("🎵 Generating audio..."):
                        language_code = converter.supported_languages[selected_language]['code']
                        audio_bytes, mime_type = converter.text_to_speech(input_text, language_code)
                        if audio_bytes and mime_type:
                            st.session_state.audio_data = {
                                'bytes': audio_bytes,
                                'format': mime_type
                            }
                            st.success("✨ Audio generated successfully!")
                
                if 'audio_data' in st.session_state:
                    # Custom CSS to make audio controls larger and more accessible
                    st.markdown("""
                        <style>
                        audio {
                            width: 100% !important;
                            height: 60px !important;
                            border-radius: 10px;
                            background-color: rgba(255, 255, 255, 0.05);
                            padding: 5px;
                        }
                        audio::-webkit-media-controls-panel {
                            background-color: rgba(255, 255, 255, 0.1);
                        }
                        </style>
                    """, unsafe_allow_html=True)
                    st.audio(
                        st.session_state.audio_data['bytes'],
                        format=st.session_state.audio_data['format']
                    )
    
    with file_tab:
        st.markdown("### 📄 Upload Text File")
        st.caption("Upload a text file in any supported language")
        
        uploaded_file = st.file_uploader(
            "",  # Hide label as we have a caption
            type=["txt"],
            help="File should be in plain text format (.txt)"
        )
        
        if uploaded_file is not None:
            try:
                # Read and detect file encoding
                file_content = uploaded_file.read()
                detected_encoding = chardet.detect(file_content)['encoding'] or 'utf-8'
                text_content = file_content.decode(detected_encoding).strip()
                
                st.markdown("---")
                
                # Show original text
                st.markdown("### 📝 Original Text")
                st.caption(f"File: {uploaded_file.name}")
                st.text_area(
                    "",  # Hide label
                    value=text_content,
                    height=150,
                    disabled=True,
                    key="original_content"
                )
                
                st.markdown("---")
                
                # Add Summary section before Bionic Text
                st.markdown("### 📋 Summary")
                summary_col1, summary_col2 = st.columns([1, 3])
                
                with summary_col1:
                    if st.button("📝 Generate Summary", key="file_summary", use_container_width=True, help="Generate a summary of the text"):
                        with st.spinner("Generating summary..."):
                            summary = converter.summarize_text(text_content, selected_language)
                            if summary:
                                st.session_state.file_text_summary = summary
                                st.success("✨ Summary generated!")
                
                if 'file_text_summary' in st.session_state:
                    with summary_col2:
                        st.markdown(
                            f"""
                            <div style='background-color: {settings["background_color"]}; 
                                      padding: 15px; 
                                      border-radius: 10px;
                                      font-size: {max(12, settings["font_size"]-4)}px;
                                      color: {"white" if theme == "Dark" else "black"};
                                      box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                                      margin: 5px 0;'>
                                {st.session_state.file_text_summary}
                            </div>
                            """,
                            unsafe_allow_html=True
                        )
                
                st.markdown("---")
                
                # Process and show bionic text
                st.markdown("### 🔍 Bionic Text")
                processed_text = converter.process_text(text_content, settings)
                
                st.markdown(
                    f"""
                    <div style='background-color: {settings["background_color"]}; 
                              padding: 20px; 
                              border-radius: 10px;
                              font-size: {settings["font_size"]}px;
                              color: {"white" if theme == "Dark" else "black"};
                              box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                              margin: 10px 0;'>
                        {processed_text}
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                
                # Export to Word option
                st.markdown("---")
                st.markdown("### 📄 Export Options")
                export_col1, export_col2, export_col3 = st.columns([2, 4, 2])
                with export_col1:
                    if st.button("📎 Export to Word", key="export_file", use_container_width=True, help="Download as Word document (.docx)"):
                        try:
                            doc_bytes = converter.export_to_word(text_content, settings)
                            with export_col2:
                                st.download_button(
                                    label="⬇️ Download Word Document",
                                    data=doc_bytes,
                                    file_name=f"bionic_text_{settings['language'].lower()}_{uploaded_file.name.rsplit('.', 1)[0]}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                        except Exception as e:
                            st.error(f"Error creating Word document: {str(e)}")
                
                # Text to Speech section
                st.markdown("---")
                st.markdown("### 🔊 Text to Speech")
                st.caption("Listen to the original text in its native language")
                
                cols = st.columns([2, 4, 2])
                with cols[0]:
                    tts_button = st.button(
                        "🎙️ Generate Audio",
                        key="tts_file",
                        use_container_width=True,
                        help="Press Space or Enter when focused to generate audio"
                    )
                
                with cols[1]:
                    if tts_button:
                        with st.spinner("🎵 Generating audio..."):
                            language_code = converter.supported_languages[selected_language]['code']
                            audio_bytes, mime_type = converter.text_to_speech(text_content, language_code)
                            if audio_bytes and mime_type:
                                st.session_state.file_audio_data = {
                                    'bytes': audio_bytes,
                                    'format': mime_type
                                }
                                st.success("✨ Audio generated successfully!")
                    
                    if 'file_audio_data' in st.session_state:
                        # Custom CSS to make audio controls larger
                        st.markdown("""
                            <style>
                            audio {
                                width: 100% !important;
                                height: 60px !important;
                                border-radius: 10px;
                                background-color: rgba(255, 255, 255, 0.05);
                                padding: 5px;
                            }
                            audio::-webkit-media-controls-panel {
                                background-color: rgba(255, 255, 255, 0.1);
                            }
                            </style>
                        """, unsafe_allow_html=True)
                        st.audio(
                            st.session_state.file_audio_data['bytes'],
                            format=st.session_state.file_audio_data['format']
                        )
            
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
                st.info("Please ensure the file is a valid text file in one of the supported languages.")

if __name__ == "__main__":
    main()
